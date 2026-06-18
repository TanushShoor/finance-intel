import docx as pydocx
from app.ingestion.base import ParsedDocument


class DocxParser:
    def parse(self, path: str) -> ParsedDocument:
        d = pydocx.Document(path)
        blocks = [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]
        return ParsedDocument(blocks=blocks, page_count=1,
                              had_text_layer=True, source=path)
