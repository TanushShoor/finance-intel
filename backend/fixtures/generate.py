import json
import os
import docx as pydocx
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

HERE = os.path.dirname(__file__)
OUT = os.path.join(HERE, "generated")
KEYS = os.path.join(HERE, "answer_keys")

STANDARD_CLAUSES = {
    "indemnity": "8. Indemnity. Each party shall indemnify the other against third-party "
                 "claims arising from its own breach of this Agreement, subject to the "
                 "limitations in Section 9.",
    "limitation_of_liability": "9. Limitation of Liability. Except for breaches of "
                 "confidentiality, each party's aggregate liability shall not exceed the fees "
                 "paid in the preceding 12 months. Neither party is liable for indirect or "
                 "consequential loss.",
    "governing_law": "10. Governing Law. This Agreement is governed by the laws of England "
                 "and Wales.",
    "termination": "11. Termination. Either party may terminate on 30 days' written notice, "
                 "or immediately for material breach not cured within 14 days.",
    "ip_ownership": "12. Intellectual Property. Each party retains ownership of its "
                 "pre-existing IP. Deliverables created for the Customer are assigned to the "
                 "Customer on payment.",
    "payment_terms": "13. Payment. The Customer shall pay undisputed invoices within 30 days "
                 "of receipt. Late amounts accrue interest at 2% per annum.",
    "confidentiality": "14. Confidentiality. Each party shall protect the other's Confidential "
                 "Information and use it only for the purposes of this Agreement, for 3 years "
                 "after termination.",
}

# Deliberately problematic versions (planted risk) + the reason each is risky.
PLANTED = {
    "indemnity": ("8. Indemnity. The Customer shall indemnify, defend and hold harmless the "
                  "Supplier from any and all claims of any kind whatsoever, including the "
                  "Supplier's own negligence, without limitation.",
                  "One-sided, uncapped indemnity covering the other party's own negligence."),
    "limitation_of_liability": ("9. Limitation of Liability. The Supplier's liability is "
                  "excluded entirely and in no event shall the Supplier be liable for any "
                  "amount. The Customer waives all claims.",
                  "Supplier liability fully excluded; customer waives all remedies."),
    "termination": ("11. Termination. The Supplier may terminate at any time for any reason "
                  "with no notice. This Agreement auto-renews for successive 5-year terms "
                  "unless cancelled 180 days in advance.",
                  "Unilateral no-notice termination plus long auto-renewal lock-in."),
    "ip_ownership": ("12. Intellectual Property. All intellectual property created by either "
                  "party, including the Customer's pre-existing IP, is assigned exclusively to "
                  "the Supplier.",
                  "Customer's pre-existing IP assigned away to the Supplier."),
    "payment_terms": ("13. Payment. All fees are due immediately on invoice and are "
                  "non-refundable. The Supplier may increase fees at any time without notice.",
                  "Immediate non-refundable payment with unilateral price increases."),
    "governing_law": ("10. Governing Law. This Agreement is governed by the laws of a "
                  "jurisdiction to be nominated by the Supplier at the time of any dispute.",
                  "Governing law unilaterally chosen by one party at dispute time."),
    "confidentiality": ("14. Confidentiality. The Customer's confidential information may be "
                  "disclosed by the Supplier to any third party at the Supplier's discretion.",
                  "Confidentiality obligation effectively removed for one party."),
}


def _write_docx(path, intro, clauses: dict):
    d = pydocx.Document()
    d.add_paragraph("MASTER SERVICES AGREEMENT")
    d.add_paragraph(intro)
    for text in clauses.values():
        d.add_paragraph(text)
    d.save(path)


def _write_pdf(path, title, clauses: dict):
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    y = height - 72
    c.setFont("Helvetica-Bold", 12); c.drawString(72, y, title); y -= 24
    c.setFont("Helvetica", 9)
    for text in clauses.values():
        for line in _wrap(text, 95):
            if y < 72:
                c.showPage(); c.setFont("Helvetica", 9); y = height - 72
            c.drawString(72, y, line); y -= 13
        y -= 8
    c.save()


def _wrap(text, n):
    words, line, out = text.split(), "", []
    for w in words:
        if len(line) + len(w) + 1 > n:
            out.append(line); line = w
        else:
            line = f"{line} {w}".strip()
    if line: out.append(line)
    return out


def main():
    os.makedirs(OUT, exist_ok=True)
    os.makedirs(os.path.join(OUT, "batch"), exist_ok=True)
    os.makedirs(KEYS, exist_ok=True)

    # Fixture 1: all clause types present (metric 1)
    _write_docx(os.path.join(OUT, "all_clauses.docx"),
                "This agreement contains the standard commercial terms.", STANDARD_CLAUSES)
    json.dump({"expected_present": list(STANDARD_CLAUSES.keys())},
              open(os.path.join(KEYS, "all_clauses.json"), "w"), indent=2)

    # Fixture 2: planted risky clauses (metrics 2 & 4)
    planted_clauses = {k: v[0] for k, v in PLANTED.items()}
    _write_pdf(os.path.join(OUT, "planted_risk.pdf"), "SERVICES AGREEMENT", planted_clauses)
    json.dump({"expected_flagged": list(PLANTED.keys()),
               "reasons": {k: v[1] for k, v in PLANTED.items()}},
              open(os.path.join(KEYS, "planted_risk.json"), "w"), indent=2)

    # Fixture 3: 3-contract batch differing on governing_law + limitation_of_liability (metric 5)
    variants = {
        "contract_a": {**STANDARD_CLAUSES,
            "governing_law": "10. Governing Law. Governed by the laws of England and Wales.",
            "limitation_of_liability": "9. Limitation of Liability. Liability capped at 12 months' fees."},
        "contract_b": {**STANDARD_CLAUSES,
            "governing_law": "10. Governing Law. Governed by the laws of the State of New York.",
            "limitation_of_liability": "9. Limitation of Liability. Liability capped at 3 months' fees."},
        "contract_c": {**STANDARD_CLAUSES,
            "governing_law": "10. Governing Law. Governed by the laws of Singapore.",
            "limitation_of_liability": "9. Limitation of Liability. Liability uncapped."},
    }
    for name, clauses in variants.items():
        _write_docx(os.path.join(OUT, "batch", f"{name}.docx"),
                    "Batch comparison fixture.", clauses)
    json.dump({"clause_type": "governing_law",
               "expected_differences": ["England and Wales", "New York", "Singapore"]},
              open(os.path.join(KEYS, "batch.json"), "w"), indent=2)
    print("Fixtures written to", OUT)


if __name__ == "__main__":
    main()
