import os
from app.ingestion.base import ParsedDocument


def test_parsed_document_full_text_joins_blocks():
    doc = ParsedDocument(
        blocks=["This is the first paragraph of the contract.",
                "This is the second paragraph of the contract."],
        page_count=1, had_text_layer=True, source="x.pdf")
    assert doc.full_text == ("This is the first paragraph of the contract.\n"
                             "This is the second paragraph of the contract.")
    assert doc.is_degraded() is False


def test_parsed_document_degraded_when_almost_no_text():
    doc = ParsedDocument(blocks=[""], page_count=3, had_text_layer=False, source="x.pdf")
    assert doc.is_degraded() is True


import docx as pydocx
from app.ingestion.docx import DocxParser


def test_docx_parser_extracts_paragraphs(tmp_path):
    p = tmp_path / "t.docx"
    d = pydocx.Document()
    d.add_paragraph("1. Indemnity")
    d.add_paragraph("The Supplier shall indemnify the Customer.")
    d.save(p)
    out = DocxParser().parse(str(p))
    assert "Indemnity" in out.full_text
    assert out.had_text_layer is True


from reportlab.pdfgen import canvas
from app.ingestion.pdf import PdfParser


def test_pdf_parser_extracts_text(tmp_path):
    p = tmp_path / "t.pdf"
    c = canvas.Canvas(str(p))
    c.drawString(72, 720, "1. Governing Law")
    c.drawString(72, 700, "This Agreement is governed by the laws of England.")
    c.save()
    out = PdfParser().parse(str(p))
    assert "Governing Law" in out.full_text
    assert out.had_text_layer is True


from app.ingestion.docling_parser import DoclingParser


class _FakeDoc:
    def export_to_markdown(self):
        return "1. Indemnity\nThe Supplier shall indemnify the Customer."


class _FakeResult:
    document = _FakeDoc()


class _FakeConverter:
    def convert(self, path):
        return _FakeResult()


def test_docling_parser_uses_converter(tmp_path):
    p = tmp_path / "scan.pdf"
    p.write_bytes(b"%PDF-1.4")
    out = DoclingParser(converter=_FakeConverter()).parse(str(p))
    assert "Indemnity" in out.full_text
    assert out.had_text_layer is True


from app.ingestion.router import parse_document


class _StubParser:
    def __init__(self, doc): self.doc = doc
    def parse(self, path): return self.doc


def test_router_falls_back_when_degraded():
    degraded = ParsedDocument(blocks=[""], page_count=2, had_text_layer=False, source="s.pdf")
    good = ParsedDocument(blocks=["Real contract text here, plenty of it."],
                          page_count=2, had_text_layer=True, source="s.pdf")
    out = parse_document("s.pdf", pdf_parser=_StubParser(degraded),
                         docx_parser=_StubParser(degraded),
                         docling_parser=_StubParser(good))
    assert out.full_text.startswith("Real contract")


def test_router_keeps_primary_when_good():
    good = ParsedDocument(blocks=["Plenty of real contract text present here."],
                          page_count=1, had_text_layer=True, source="s.docx")
    sentinel = ParsedDocument(blocks=["FALLBACK"], page_count=1, had_text_layer=True, source="s")
    out = parse_document("s.docx", pdf_parser=_StubParser(good),
                         docx_parser=_StubParser(good),
                         docling_parser=_StubParser(sentinel))
    assert "FALLBACK" not in out.full_text


FIX = os.path.join(os.path.dirname(__file__), "..", "fixtures", "generated")


def test_fixture_all_clauses_parses():
    out = parse_document(os.path.join(FIX, "all_clauses.docx"))
    assert "Indemnity" in out.full_text and "Confidentiality" in out.full_text


def test_fixture_planted_risk_pdf_parses():
    out = parse_document(os.path.join(FIX, "planted_risk.pdf"))
    assert "indemnify" in out.full_text.lower()
