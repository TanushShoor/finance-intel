from app.pipeline.runner import run_financial_analysis
from app.schemas.models import ChunkBlocks, DocumentOutline
from app.schemas.financial import (
    DocumentIdentity, MetricSet, Metric, ToneAnalysis, RiskFactorSet, RiskFactor,
    InvestmentMemo)


def test_run_financial_analysis_assembles_all_stages(mock_llm, tmp_path):
    # Order matches the runner: structure(map, synth) → identity → metrics → tone
    # → risk factors → memo. Short fixture => 1 chunk per mapped stage.
    mock_llm.queue_response(ChunkBlocks(blocks=[]))
    mock_llm.queue_response(DocumentOutline(title="Intel Q4", outline=[]))
    mock_llm.queue_response(DocumentIdentity(company="Intel", period="Q4 2024",
                                             doc_type="earnings release"))
    mock_llm.queue_response(MetricSet(metrics=[
        Metric(name="revenue", value="$14.3B", period="Q4 2024")]))
    mock_llm.queue_response(ToneAnalysis(overall_sentiment="cautious", confidence_score=40))
    mock_llm.queue_response(RiskFactorSet(risk_factors=[
        RiskFactor(category="market", title="Demand softness", severity=60)]))
    mock_llm.queue_response(InvestmentMemo(company_overview="Intel.",
                                           bull_case=["upside"], bear_case=["downside"]))

    fixture = tmp_path / "intel.docx"
    import docx as pydocx
    d = pydocx.Document()
    d.add_paragraph("Intel reported Q4 revenue of $14.3 billion.")
    d.save(fixture)

    a = run_financial_analysis(str(fixture), llm=mock_llm)
    assert a.identity.company == "Intel"
    assert a.metrics[0].name == "revenue"
    assert a.tone.overall_sentiment == "cautious"
    assert a.risk_factors[0].title == "Demand softness"
    assert a.memo.bull_case == ["upside"] and a.memo.bear_case == ["downside"]


def test_one_failing_stage_does_not_sink_run(mock_llm, tmp_path):
    # structure ok, then every later call raises -> stages fall back to defaults,
    # but the run still returns a populated structure and empty analyses.
    mock_llm.queue_response(ChunkBlocks(blocks=[]))
    mock_llm.queue_response(DocumentOutline(title="X", outline=[]))
    # No more responses queued: MockLLM raises AssertionError on the next call,
    # which each stage's try/except swallows.

    fixture = tmp_path / "f.docx"
    import docx as pydocx
    d = pydocx.Document(); d.add_paragraph("Some short text."); d.save(fixture)

    a = run_financial_analysis(str(fixture), llm=mock_llm)
    assert a.structure.title == "X"
    assert a.metrics == [] and a.risk_factors == []
    assert a.memo.bull_case == []
